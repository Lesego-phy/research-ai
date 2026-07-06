# PDF / Word / Markdown Export Tools

import os
from fpdf import FPDF
from docx import Document
from datetime import datetime

# Ensure the exports directory exists
EXPORT_DIR = "data/exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

def generate_filename(topic: str, extension: str) -> str:
    """Creates a safe, timestamped filename."""
    safe_topic = "".join(c for c in topic if c.isalnum() or c in (' ', '-')).rstrip()[:30]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(EXPORT_DIR, f"{safe_topic}_{timestamp}.{extension}")

def export_to_markdown(text: str, topic: str) -> str:
    filepath = generate_filename(topic, "md")
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(text)
    return filepath

def export_to_pdf(text: str, topic: str) -> str:
    filepath = generate_filename(topic, "pdf")
    pdf = FPDF()
    pdf.add_page()
    # Set font (Arial, normal, 12pt)
    pdf.set_font("Arial", size=12)
    
    # FPDF requires latin-1 encoding by default, so we encode/decode to handle special chars
    # Multi_cell handles word wrapping automatically
    pdf.multi_cell(0, 10, text.encode('latin-1', 'replace').decode('latin-1'))
    
    pdf.output(filepath)
    return filepath

def export_to_docx(text: str, topic: str) -> str:
    filepath = generate_filename(topic, "docx")
    doc = Document()
    
    # Add Title
    doc.add_heading(topic, 0)
    
    # Add Content
    # Split by paragraphs to preserve formatting
    for paragraph in text.split('\n'):
        doc.add_paragraph(paragraph)
        
    doc.save(filepath)
    return filepath